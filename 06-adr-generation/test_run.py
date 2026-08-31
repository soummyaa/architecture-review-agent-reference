import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import MagicMock, patch

from docx import Document

from run import (
    DEFAULT_TEMPLATE_PATH,
    AdrCondition,
    ArchitectureDecisionRecord,
    Citation,
    SharePointWriteError,
    load_adr,
    render_adr,
    upload_to_sharepoint,
)


def example_adr() -> ArchitectureDecisionRecord:
    citation = Citation(
        standard_id="STD-002",
        section="2. Workload authentication",
        source_file="STD-002-identity-and-access.md",
    )
    return ArchitectureDecisionRecord(
        submission_id="SUB-001",
        technology="Northwind Analytics Cloud",
        title="Adopt Northwind Analytics Cloud",
        status="proposed",
        decision="approved_with_conditions",
        context="The analytics platform is under architecture review.",
        standards_assessment="The review identified a workload identity gap.",
        decision_statement="Approve after the identity gap is addressed.",
        decision_drivers=["Conformance with enterprise identity standards."],
        conditions=[
            AdrCondition(
                action="Replace the static credential with an approved identity pattern.",
                rationale="Static credentials do not meet the standard.",
                citation=citation,
            )
        ],
        positive_consequences=["The analytics capability can proceed."],
        negative_consequences=["Identity integration requires additional work."],
    )


def document_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join(paragraphs + table_cells)


class LocalRenderingTests(unittest.TestCase):
    def test_loads_reviewed_adr_from_workflow_output(self) -> None:
        payload = {
            "standards_report": {},
            "draft_adr": {},
            "review": {"reviewed_adr": example_adr().model_dump(mode="json")},
        }
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "reviewed.json"
            path.write_text(json.dumps(payload), encoding="utf-8")

            adr = load_adr(path)

        self.assertEqual(adr.submission_id, "SUB-001")

    def test_renders_structured_adr_into_template(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output_path = Path(directory) / "SUB-001-adr.docx"

            rendered_path = render_adr(example_adr(), DEFAULT_TEMPLATE_PATH, output_path)

            text = document_text(rendered_path)
        self.assertIn("Adopt Northwind Analytics Cloud", text)
        self.assertIn("Replace the static credential", text)
        self.assertIn("STD-002, 2. Workload authentication", text)
        self.assertNotIn("{{", text)
        self.assertNotIn("{%", text)


class SharePointUploadTests(unittest.TestCase):
    @patch("run.requests.Session")
    @patch("run.DefaultAzureCredential")
    def test_reports_missing_graph_write_permission_clearly(
        self,
        credential_factory: MagicMock,
        session_factory: MagicMock,
    ) -> None:
        credential = credential_factory.return_value.__enter__.return_value
        credential.get_token.return_value.token = "token"
        session = session_factory.return_value.__enter__.return_value

        site_response = MagicMock(ok=True)
        site_response.json.return_value = {"id": "example-site-id"}
        session.get.return_value = site_response
        upload_response = MagicMock(ok=False, status_code=403, reason="Forbidden")
        upload_response.headers = {"request-id": "request-123"}
        session.put.return_value = upload_response

        with tempfile.TemporaryDirectory() as directory:
            document_path = Path(directory) / "SUB-001-adr.docx"
            document_path.write_bytes(b"local document")

            with self.assertRaisesRegex(
                SharePointWriteError,
                "denied permission.*Graph permission that allows writes",
            ):
                upload_to_sharepoint(
                    document_path,
                    "example.sharepoint.com",
                    "/sites/architecture-workshop",
                    "Architecture Reviews",
                )

            self.assertTrue(document_path.is_file())


if __name__ == "__main__":
    unittest.main()