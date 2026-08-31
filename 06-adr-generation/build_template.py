#!/usr/bin/env python3
"""Build the workshop's editable DOCX template."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

MODULE_DIRECTORY = Path(__file__).resolve().parent
TEMPLATE_PATH = MODULE_DIRECTORY / "adr-template.docx"


def build_template(output_path: Path = TEMPLATE_PATH) -> Path:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(26)

    label = document.add_paragraph("ARCHITECTURE DECISION RECORD")
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label.runs[0].bold = True
    label.runs[0].font.size = Pt(9)
    document.add_paragraph("{{ adr.title }}", style="Title").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    metadata = document.add_table(rows=4, cols=2)
    metadata.style = "Light Shading Accent 1"
    metadata_values = (
        ("Submission", "{{ adr.submission_id }}"),
        ("Technology", "{{ adr.technology }}"),
        ("ADR status", "{{ adr.status | replace('_', ' ') | title }}"),
        ("Decision", "{{ adr.decision | replace('_', ' ') | title }}"),
    )
    for row, (heading, value) in zip(metadata.rows, metadata_values, strict=True):
        row.cells[0].text = heading
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = value
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    document.add_heading("Context", level=1)
    document.add_paragraph("{{ adr.context }}")
    document.add_heading("Standards assessment", level=1)
    document.add_paragraph("{{ adr.standards_assessment }}")
    document.add_heading("Decision", level=1)
    document.add_paragraph("{{ adr.decision_statement }}")

    document.add_heading("Decision drivers", level=1)
    document.add_paragraph("{%p for driver in adr.decision_drivers %}")
    document.add_paragraph("{{ driver }}", style="List Bullet")
    document.add_paragraph("{%p endfor %}")

    document.add_heading("Conditions", level=1)
    document.add_paragraph("{%p if adr.conditions %}")
    document.add_paragraph("{%p for condition in adr.conditions %}")
    document.add_paragraph("Condition {{ loop.index }}", style="Heading 2")
    document.add_paragraph("{{ condition.action }}")
    rationale = document.add_paragraph()
    rationale.add_run("Rationale: ").bold = True
    rationale.add_run("{{ condition.rationale }}")
    source = document.add_paragraph()
    source.add_run("Source: ").bold = True
    source.add_run(
        "{{ condition.citation.standard_id }}, {{ condition.citation.section }} "
        "({{ condition.citation.source_file }})"
    )
    document.add_paragraph("{%p endfor %}")
    document.add_paragraph("{%p else %}")
    document.add_paragraph("None.")
    document.add_paragraph("{%p endif %}")

    document.add_heading("Positive consequences", level=1)
    document.add_paragraph("{%p for consequence in adr.positive_consequences %}")
    document.add_paragraph("{{ consequence }}", style="List Bullet")
    document.add_paragraph("{%p endfor %}")

    document.add_heading("Negative consequences", level=1)
    document.add_paragraph("{%p for consequence in adr.negative_consequences %}")
    document.add_paragraph("{{ consequence }}", style="List Bullet")
    document.add_paragraph("{%p endfor %}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


if __name__ == "__main__":
    print(build_template())